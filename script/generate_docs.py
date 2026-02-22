#!/usr/bin/env python3
"""
generate_docs.py — Islamic Will Word Document Generator

Generates formatted .docx documents from a JSON data file
produced by the Islamic Will skill (/islamic-will).

Usage:
    python script/generate_docs.py will_data.json
    python script/generate_docs.py will_data.json --output-dir ./output

Install dependency:
    uv add python-docx          # if using uv
    pip install python-docx     # if using pip
"""

import json
import sys
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ─────────────────────────────────────────────────────────────────
FONT      = "Times New Roman"
BODY_PT   = 12
TITLE_PT  = 16
HEAD_PT   = 13
SMALL_PT  = 10
MARGIN    = Inches(1.25)


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_font(run, size=BODY_PT, bold=False, italic=False):
    run.font.name  = FONT
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic


def _para(doc, text="", size=BODY_PT, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=6,
          left_indent=None):
    """Create a paragraph with consistent font styling."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic)
    return p


def rule(doc):
    """Horizontal rule via paragraph bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "6")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "000000")
    pBdr.append(btm)
    pPr.append(pBdr)
    return p


def doc_title(doc, text):
    return _para(doc, text, size=TITLE_PT, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)


def sub_title(doc, text):
    return _para(doc, text, size=BODY_PT,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)


def article(doc, text):
    p = _para(doc, text, size=HEAD_PT, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER,
               space_before=12, space_after=6)
    return p


def body(doc, text, bold=False, italic=False, size=BODY_PT,
         indent=None, space_before=0, space_after=6):
    return _para(doc, text, size=size, bold=bold, italic=italic,
                 left_indent=indent,
                 space_before=space_before, space_after=space_after)


def sig_block(doc, label, date_line=True, name_label=None, indent=0):
    """Signature line + optional date + printed name."""
    p = _para(doc, "_" * 48, space_before=24, space_after=2,
              left_indent=indent if indent else None)
    lbl = _para(doc, label, size=BODY_PT, space_before=2, space_after=2,
                left_indent=indent if indent else None)
    if date_line:
        _para(doc, "Date: " + "_" * 20, space_before=4, space_after=8,
              left_indent=indent if indent else None)
    return p


def detail_line(doc, label, indent=0):
    _para(doc, f"{label}: " + "_" * 38,
          space_before=2, space_after=2,
          left_indent=indent if indent else None)


def _new_doc() -> Document:
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin   = MARGIN
        s.right_margin  = MARGIN
    return doc


def _safe(name: str) -> str:
    return name.replace(" ", "_").replace("/", "_")


def _table_header(table, headers, size=BODY_PT):
    cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cells[i].text = h
        for run in cells[i].paragraphs[0].runs:
            _set_font(run, size=size, bold=True)


def _table_row(table, values, bold=False, size=BODY_PT):
    cells = table.add_row().cells
    for i, v in enumerate(values):
        cells[i].text = str(v)
        for run in cells[i].paragraphs[0].runs:
            _set_font(run, size=size, bold=bold)
    return cells


def _disclaimer(doc):
    rule(doc)
    _para(doc,
          "This document was prepared with AI assistance and must be reviewed by a "
          "Tennessee-licensed estate planning attorney before execution.",
          size=SMALL_PT, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)


# ── Generator class ───────────────────────────────────────────────────────────

class IslamicWillGenerator:

    def __init__(self, data: dict, output_dir: Path):
        self.data   = data
        self.out    = output_dir
        self.out.mkdir(parents=True, exist_ok=True)

        self.t    = data.get("testator", {})
        self.sp   = data.get("spouse", {})
        self.h    = data.get("heirs", {})
        self.est  = data.get("estate", {})
        self.was  = data.get("wasiyyah", [])
        self.exe  = data.get("executor", {})
        self.grd  = data.get("guardian", {})
        self.dist = data.get("distribution", {})

    # ── helpers ───────────────────────────────────────────────────────────────

    def _name(self):    return self.t.get("full_name", "[TESTATOR NAME]")
    def _county(self):  return self.t.get("county",    "[COUNTY]")
    def _madhhab(self): return self.t.get("madhhab",   "Hanafi")
    def _dob(self):     return self.t.get("dob",        "[DOB]")

    def _elective_pct(self):
        y = self.sp.get("years_married", 0)
        if y < 3:  return "10%"
        if y < 6:  return "20%"
        if y < 9:  return "30%"
        return "40%"

    def _spouse_share(self):
        for s in self.dist.get("shares", []):
            if s.get("relationship") == "Wife":
                return s
        return {}

    def _wasiyyah_total(self):
        return sum(w.get("amount", 0) for w in self.was)

    # ── generate all ──────────────────────────────────────────────────────────

    def generate_all(self):
        steps = [
            ("Last Will and Testament",       self.gen_will),
            ("Waiver of Elective Share",       self.gen_waiver),
            ("Execution Checklist",            self.gen_checklist),
            ("Beneficiary Designation Reminder", self.gen_beneficiary),
        ]
        paths = []
        for i, (label, fn) in enumerate(steps, 1):
            print(f"  [{i}/{len(steps)}] {label}...")
            paths.append(fn())
        return paths

    # ── DOCUMENT 1 — Last Will and Testament ─────────────────────────────────

    def gen_will(self) -> Path:
        doc  = _new_doc()
        name    = self._name()
        county  = self._county()
        madhhab = self._madhhab()
        dob     = self._dob()
        dist    = self.dist
        shares  = dist.get("shares", [])
        net     = dist.get("net_mirath", 0)

        # ── Title ──
        doc_title(doc, "LAST WILL AND TESTAMENT")
        doc_title(doc, "OF")
        doc_title(doc, name.upper())
        sub_title(doc, "STATE OF TENNESSEE")
        sub_title(doc, f"COUNTY OF {county.upper().replace(' COUNTY','')}")
        rule(doc)

        body(doc,
             f"I, {name.upper()}, a resident of {county}, Tennessee, born on {dob}, "
             "being of sound mind and disposing memory, and not acting under duress, "
             "menace, fraud, or undue influence of any person whomsoever, do hereby "
             "make, publish, and declare this instrument to be my Last Will and "
             "Testament, hereby revoking all former wills and codicils previously "
             "made by me.", space_after=12)

        # ── Article I — Religious Declaration ──
        rule(doc); article(doc, "ARTICLE I — RELIGIOUS DECLARATION")
        body(doc,
             f"I am a Muslim of the Sunni faith, following the {madhhab} school of "
             f"jurisprudence (Madhhab). It is my sincere intention and express "
             f"direction that my estate be distributed in accordance with the Islamic "
             f"Law of Inheritance (Ilm al-Fara'id) as interpreted by the {madhhab} "
             "school. This will has been drafted to honor that intent while complying "
             "with the laws of the State of Tennessee.")

        # ── Article II — Personal Information ──
        rule(doc); article(doc, "ARTICLE II — PERSONAL INFORMATION")
        body(doc, f"I was born on {dob} and reside in {county}, Tennessee.")

        # ── Article III — Executor ──
        rule(doc); article(doc, "ARTICLE III — APPOINTMENT OF EXECUTOR (WASI)")
        ep = self.exe.get("primary",   "[PRIMARY EXECUTOR]")
        ea = self.exe.get("alternate", "[ALTERNATE EXECUTOR]")
        for clause, text in [
            ("3.1  Primary Executor",
             f"I hereby appoint {ep.upper()} as the Executor (Wasi) of this Will. "
             "It is my preference that my Executor be a Muslim of good character who "
             "understands and respects my religious obligations."),
            ("3.2  Alternate Executor",
             f"If {ep} is unable or unwilling to serve, I appoint "
             f"{ea.upper()} as alternate Executor."),
            ("3.3  Powers of Executor",
             "My Executor shall have full power to: collect and manage all estate "
             "assets; pay all funeral expenses, religious obligations, debts, taxes, "
             "and administrative expenses; sell, mortgage, or lease real or personal "
             "property as necessary; and distribute the estate as directed herein, "
             "without being required to obtain court approval for individual acts, to "
             "the extent permitted by Tennessee law."),
            ("3.4  Bond",
             "I waive the requirement that my Executor post a bond."),
        ]:
            body(doc, clause, bold=True, space_before=6, space_after=2)
            body(doc, text, indent=0.3)

        # ── Article IV — Funeral ──
        rule(doc); article(doc, "ARTICLE IV — FUNERAL AND BURIAL INSTRUCTIONS (TAJHEEZ WA TAKFEEN)")
        body(doc,
             "I direct that my funeral and burial be conducted in strict accordance "
             "with Sunni Islamic tradition:")
        funeral_items = [
            "(a)  Ghusl: My body shall be washed by a qualified Muslim of the same gender, per the Sunnah;",
            "(b)  Kafan: My body shall be wrapped in white unsewn linen cloth (Kafan);",
            "(c)  Janazah Prayer: An Islamic funeral prayer (Salat al-Janazah) shall be performed;",
            f"(d)  Burial: I shall be buried in a Muslim cemetery in or near {county}, Tennessee, facing the Qiblah if possible;",
            "(e)  No Embalming: My body shall NOT be embalmed unless required by law;",
            "(f)  No Open Casket: There shall be no open casket viewing;",
            "(g)  No Cremation: I expressly prohibit cremation of my body.",
        ]
        for item in funeral_items:
            body(doc, item, indent=0.3, space_after=3)
        funeral = self.est.get("funeral_expenses", 5000)
        body(doc,
             f"Reasonable funeral and burial expenses, not to exceed ${funeral:,.0f}, "
             "shall be paid from my estate as the first priority before any distribution.")

        # ── Article V — Debts ──
        rule(doc); article(doc, "ARTICLE V — PAYMENT OF DEBTS (DUYUN)")
        body(doc,
             "I direct my Executor to pay all of my lawful debts from my estate "
             "before any inheritance distribution:")
        body(doc, "5.1  Religious Obligations:", bold=True, space_after=2)
        for item in [
            "(a)  Any unpaid Mahr (dowry) owed to my wife;",
            "(b)  Any unpaid Zakat (obligatory charity) from prior years;",
            "(c)  Any Kaffarah (religious expiation) owed;",
            "(d)  Any cost of Hajj by proxy (Hajj al-Badal) if applicable.",
        ]:
            body(doc, item, indent=0.6, space_after=3)
        body(doc, "5.2  Secular Debts:", bold=True, space_after=2)
        body(doc,
             "All outstanding mortgages, loans, credit cards, medical bills, and "
             "other lawful debts proven by creditors within the time permitted by "
             "Tennessee law.", indent=0.3)

        # ── Article VI — Wasiyyah ──
        if self.was:
            rule(doc); article(doc, "ARTICLE VI — BEQUEST (WASIYYAH)")
            net_debts = self.est.get("probate", 0) - self.est.get("funeral_expenses", 0)
            body(doc,
                 f"6.1  Limitation: These bequests shall not exceed one-third (1/3) "
                 f"of the net estate (maximum ${net_debts/3:,.0f}) after payment of "
                 "funeral expenses and debts, per Islamic Law.")
            body(doc, "6.2  Specific Bequests:", bold=True, space_after=2)
            for i, w in enumerate(self.was, 1):
                body(doc,
                     f"({'abcdefg'[i-1]})  To {w.get('recipient','[RECIPIENT]')}: "
                     f"${w.get('amount',0):,.0f} for the purpose of "
                     f"{w.get('purpose','[PURPOSE]')}.",
                     indent=0.3, space_after=4)
            body(doc,
                 "6.3  Failure of Bequest: If a designated recipient ceases to exist, "
                 "the Executor shall transfer the amount to the nearest equivalent "
                 "Islamic charitable organization.", indent=0.3)

        # ── Article VII — Mirath ──
        art_n = "VII" if self.was else "VI"
        rule(doc); article(doc, f"ARTICLE {art_n} — DISTRIBUTION OF ESTATE (MIRATH)")
        aul   = dist.get("aul_applied", False)
        denom = dist.get("aul_denominator", dist.get("original_denominator", 1))

        body(doc,
             f"7.1  Declaration of Intent\n"
             f"The remainder of my estate shall be distributed per the Islamic Law of "
             f"Inheritance (Fara'id), {madhhab} school.")
        if aul:
            body(doc,
                 f"7.2  Al-Aul Applied: Fixed shares exceeded 100%. Per {madhhab} "
                 f"jurisprudence, all shares are reduced proportionally (denominator "
                 f"adjusted to {denom}).", space_before=4)

        body(doc, "7.3  Distribution:", bold=True, space_before=8, space_after=6)
        if shares:
            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Table Grid"
            _table_header(tbl, ["Heir", "Relationship", "Fraction", "Percentage", "Est. Amount"])
            for s in shares:
                _table_row(tbl, [
                    s.get("heir", ""),
                    s.get("relationship", ""),
                    f"{s.get('fraction_num')}/{s.get('fraction_den')}",
                    f"{s.get('percentage', 0):.2f}%",
                    f"${s.get('amount', 0):,.0f}",
                ])
            _table_row(tbl, ["TOTAL", "", f"{denom}/{denom}", "100.00%", f"${net:,.0f}"], bold=True)
            doc.add_paragraph()

        excl = dist.get("exclusion_note", "")
        if excl:
            body(doc, f"7.4  Excluded Heirs: {excl}")
        body(doc,
             "7.5  Predeceased Heir: If any named heir predeceases the Testator, "
             "their share is redistributed per Hanafi Fara'id rules.")
        body(doc,
             "7.6  Non-Muslim Heirs: Any heir who is not Muslim at death shall "
             "not receive a Mirath share.")

        # ── Article VIII — Guardian ──
        art_g = "VIII" if self.was else "VII"
        daughters = self.h.get("daughters", [])
        sons      = self.h.get("sons", [])
        gp = self.grd.get("primary",   "[PRIMARY GUARDIAN]")
        ga = self.grd.get("alternate", "[ALTERNATE GUARDIAN]")
        spouse_name = self.sp.get("full_name", "my spouse")

        if gp or daughters or sons:
            rule(doc); article(doc, f"ARTICLE {art_g} — GUARDIAN OF MINOR CHILDREN (WALI)")
            children_str = ", ".join(daughters + sons) if (daughters or sons) else "my minor children"
            body(doc,
                 f"8.1  Appointment\nIf any of my children are minors at the time of my "
                 f"death, and if {spouse_name} is also deceased or unable to serve, "
                 f"I nominate {gp.upper()} as Guardian.")
            body(doc,
                 f"8.2  Alternate Guardian\nIf {gp} is unable or unwilling to serve, "
                 f"I nominate {ga.upper()} as alternate Guardian.")
            body(doc,
                 f"8.3  Islamic Upbringing\nIt is my express wish that {children_str} "
                 f"be raised in the Islamic faith (Sunni, {madhhab} tradition), receive "
                 "regular Islamic education, maintain their daily prayers, and be raised "
                 f"in a Muslim household. I request that {county} courts give full "
                 "weight to this instruction consistent with T.C.A. § 36-6-106.")
            body(doc, "8.4  Bond: I waive the requirement that the Guardian post a bond.")

        # ── No-Contest ──
        rule(doc); article(doc, "ARTICLE IX — NO-CONTEST CLAUSE (IN TERROREM)")
        body(doc,
             "If any person contests or challenges the validity of this Will through "
             "any legal proceeding, such person shall forfeit any and all rights under "
             "this Will, and their share shall pass as if they predeceased the Testator, "
             "per Hanafi Fara'id. Enforceable per T.C.A. § 32-3-119.")

        # ── Governing Law ──
        rule(doc); article(doc, "ARTICLE X — GOVERNING LAW")
        body(doc,
             f"This Will shall be governed by the laws of the State of Tennessee. "
             f"Probate shall occur in {county} Probate Court.")

        # ── Severability ──
        rule(doc); article(doc, "ARTICLE XI — SEVERABILITY")
        body(doc,
             "If any provision of this Will is found invalid or unenforceable, the "
             "remaining provisions shall remain in full force and effect.")

        # ── Testator Signature ──
        rule(doc); article(doc, "SIGNATURE OF TESTATOR")
        body(doc,
             f"IN WITNESS WHEREOF, I, {name.upper()}, declare this to be my Last "
             "Will and Testament and have signed it on the date indicated below.",
             space_after=18)
        sig_block(doc, f"{name.upper()}, Testator")
        body(doc, f"{county}, Tennessee", space_after=18)

        # ── Witnesses ──
        rule(doc); article(doc, "ATTESTATION OF WITNESSES")
        body(doc,
             f"We, the undersigned, certify that {name.upper()} signed this instrument "
             "in our presence; appeared to be of sound mind and not under duress; and "
             "that we signed as witnesses in the Testator's presence and in each "
             "other's presence. Neither of us is a beneficiary under this Will.",
             space_after=12)
        for i in (1, 2):
            body(doc, f"Witness {i}:", bold=True, space_after=2)
            sig_block(doc, f"Witness {i} Signature")
            detail_line(doc, "Printed Name")
            detail_line(doc, "Address")
            detail_line(doc, "City, State, Zip")
            doc.add_paragraph()

        # ── Self-Proving Affidavit ──
        rule(doc); article(doc, "SELF-PROVING AFFIDAVIT  (T.C.A. § 32-1-104)")
        sub_title(doc, "STATE OF TENNESSEE")
        sub_title(doc, f"COUNTY OF {county.upper().replace(' COUNTY','')}")
        body(doc,
             f"Before me, the undersigned Notary Public, personally appeared "
             f"{name.upper()}, [WITNESS 1 NAME], and [WITNESS 2 NAME]. The Testator "
             "declared that the foregoing instrument is their Last Will and Testament, "
             "willingly signed as their free and voluntary act. Each witness stated "
             "they signed in the presence of the Testator and each other.",
             space_after=18)
        for label in [f"{name.upper()}, Testator", "Witness 1", "Witness 2"]:
            sig_block(doc, label, date_line=False)
            doc.add_paragraph()
        body(doc,
             "Subscribed and sworn before me this ____ day of ____________, 20____.",
             space_before=12, space_after=18)
        sig_block(doc, "Notary Public, State of Tennessee", date_line=False)
        body(doc, f"County of {county.replace(' County','')}", space_before=2, space_after=2)
        detail_line(doc, "My Commission Expires")
        body(doc, "[NOTARY SEAL]", space_before=8)

        _disclaimer(doc)
        path = self.out / f"01_Last_Will_{_safe(name)}.docx"
        doc.save(path)
        return path

    # ── DOCUMENT 2 — Waiver of Elective Share ────────────────────────────────

    def gen_waiver(self) -> Path:
        doc         = _new_doc()
        name        = self._name()
        county      = self._county()
        spouse      = self.sp.get("full_name",     "[SPOUSE NAME]")
        years       = self.sp.get("years_married", 0)
        madhhab     = self._madhhab()
        elective    = self._elective_pct()
        ss          = self._spouse_share()
        islamic_pct = (f"{ss['fraction_num']}/{ss['fraction_den']} "
                       f"({ss['percentage']:.2f}%)" if ss else "[Islamic share per Fara'id]")
        islamic_amt = f"${ss['amount']:,.0f}" if ss else "[amount]"

        doc_title(doc, "WAIVER OF ELECTIVE SHARE AND BENEFITS")
        sub_title(doc, "(Pursuant to T.C.A. § 31-4-102)")
        sub_title(doc, "STATE OF TENNESSEE")
        sub_title(doc, f"COUNTY OF {county.upper().replace(' COUNTY','')}")
        rule(doc)

        article(doc, "1. PARTIES")
        body(doc,
             f"This Waiver is made this ____ day of ____________, 20____, by "
             f"{spouse.upper()} (hereinafter \"Undersigned\"), spouse of "
             f"{name.upper()} (hereinafter \"Testator\").")

        rule(doc); article(doc, "2. PURPOSE AND INTENT")
        body(doc,
             f"The Undersigned acknowledges that the Testator has executed a Last Will "
             f"and Testament intended to comply with Sunni Islamic Law (Shariah), "
             f"{madhhab} school.")
        body(doc,
             f"The Undersigned understands that under T.C.A. § 31-4-101, as a "
             f"surviving spouse married {years} years, she is entitled to an Elective "
             f"Share of {elective} of the Testator's estate.")
        body(doc,
             f"The Undersigned further understands that her Islamic share under "
             f"{madhhab} Fara'id is {islamic_pct} of the Net Mirath, estimated at "
             f"{islamic_amt}. It is her free and voluntary intent to waive her "
             "statutory rights to facilitate distribution per Islamic Law.")

        rule(doc); article(doc, "3. WAIVER OF RIGHTS")
        body(doc,
             "The Undersigned hereby voluntarily and irrevocably waives, releases, "
             "and relinquishes any and all rights as surviving spouse, including:")
        for item in [
            "(a)  The Right of Election against the Will (T.C.A. § 31-4-101);",
            "(b)  The Right to Year's Support (T.C.A. § 30-2-102);",
            "(c)  The Right to Exempt Property (T.C.A. § 30-2-101);",
            "(d)  The Right to Homestead (T.C.A. § 30-2-201).",
        ]:
            body(doc, item, indent=0.3, space_after=3)
        body(doc,
             f"The Undersigned agrees to accept only her Islamic inheritance share "
             f"({islamic_pct}) as stated in the Will.")

        rule(doc); article(doc, "4. FINANCIAL DISCLOSURE")
        body(doc,
             "The Undersigned acknowledges that she has been provided with, or has "
             "voluntarily waived the right to, a fair and reasonable disclosure of "
             "the Testator's property and financial obligations:")
        gross    = self.est.get("gross_total",    0)
        probate  = self.est.get("probate",         0)
        home_v   = self.est.get("home", {}).get("value", 0)
        ret      = self.est.get("retirement_401k", 0)
        funeral  = self.est.get("funeral_expenses", 0)
        was_tot  = self._wasiyyah_total()
        for line in [
            f"Gross estate approximately ${gross:,.0f}",
            f"Probate estate approximately ${probate:,.0f}",
            f"Non-probate: Home ${home_v:,.0f} (joint) · Retirement ${ret:,.0f}",
            f"Funeral expenses ${funeral:,.0f}",
            f"Wasiyyah bequest ${was_tot:,.0f}",
        ]:
            body(doc, "•  " + line, indent=0.3, space_after=3)

        rule(doc); article(doc, "5. INDEPENDENT COUNSEL")
        body(doc,
             "The Undersigned acknowledges that she has had the opportunity to consult "
             "with independent legal counsel regarding the legal consequences of this "
             "waiver and has either done so or knowingly and voluntarily declined.")

        rule(doc); article(doc, "6. BINDING EFFECT")
        body(doc,
             "This waiver shall be binding upon the Undersigned, her heirs, "
             "executors, administrators, and assigns.")

        rule(doc); article(doc, "SIGNATURE")
        body(doc,
             "I, the Undersigned, have read and understand this Waiver and sign it "
             "freely and voluntarily.", space_after=20)
        sig_block(doc, spouse.upper())

        rule(doc); article(doc, "NOTARY ACKNOWLEDGMENT")
        sub_title(doc, "STATE OF TENNESSEE")
        sub_title(doc, f"COUNTY OF {county.upper().replace(' COUNTY','')}")
        body(doc,
             f"Before me, the undersigned Notary Public, personally appeared "
             f"{spouse.upper()}, who acknowledged she executed this waiver freely "
             "and voluntarily.", space_after=20)
        sig_block(doc, "Notary Public, State of Tennessee", date_line=False)
        detail_line(doc, "My Commission Expires")
        body(doc, "[NOTARY SEAL]", space_before=8)

        _disclaimer(doc)
        path = self.out / f"02_Waiver_Elective_Share_{_safe(spouse)}.docx"
        doc.save(path)
        return path

    # ── DOCUMENT 3 — Execution Checklist ─────────────────────────────────────

    def gen_checklist(self) -> Path:
        doc    = _new_doc()
        name   = self._name()
        county = self._county()
        spouse = self.sp.get("full_name", "[SPOUSE]")
        ep     = self.exe.get("primary",   "[EXECUTOR]")
        ea     = self.exe.get("alternate", "[ALT EXECUTOR]")
        gp     = self.grd.get("primary",   "[GUARDIAN]")
        ga     = self.grd.get("alternate", "[ALT GUARDIAN]")

        doc_title(doc, "TENNESSEE ISLAMIC WILL")
        doc_title(doc, "EXECUTION CHECKLIST")
        sub_title(doc, f"{name}  ·  {county}")
        rule(doc)

        phases = [
            ("PHASE 1 — BEFORE THE SIGNING APPOINTMENT", [
                ("Legal Review", [
                    "Have the Will and Waiver reviewed by a Tennessee-licensed estate planning attorney",
                    "Fill in parents' full legal names in the distribution article of the Will",
                    "Name a specific Waqf organization for the Mosul bequest (Article VI)",
                    "Confirm daughters are minors if the guardian clause is to be active",
                ]),
                ("Deed Review  (CRITICAL — $750,000 at stake)", [
                    f"Pull the home deed from {county} Register of Deeds",
                    "Identify deed type: JTWROS (auto to spouse) vs. Tenants in Common (passes through Will)",
                    "If JTWROS: consult attorney about re-titling as Tenants in Common for Fara'id compliance",
                ]),
                ("Two Disinterested Witnesses", [
                    "Both must be adults (18+) and NOT named as beneficiaries in this Will",
                    f"Cannot be: {spouse}, Ayah Altalib, Leena Altalib, parents, {gp}, or {ga}",
                    "Both must attend the signing simultaneously",
                ]),
                ("Notary Public", [
                    f"Locate a Notary in {county} (banks, UPS Store, County Clerk's office)",
                    "Schedule signing appointment with both witnesses and Notary present",
                ]),
            ]),
            ("PHASE 2 — SIGNING DAY", [
                (f"{name} (Testator)", [
                    "Bring government-issued photo ID",
                    "Sign the Will in the presence of both witnesses and the Notary",
                    "Initial every page (recommended best practice)",
                    "Sign the Self-Proving Affidavit before the Notary",
                ]),
                ("Both Witnesses", [
                    "Watch the Testator sign",
                    "Sign in the Testator's presence and in each other's presence",
                    "Print full legal name and current address",
                ]),
                ("Notary", [
                    "Administer oath for the Self-Proving Affidavit",
                    "All three (Testator + 2 witnesses) sign affidavit before Notary",
                    "Notary affixes official seal and signs with commission expiration date",
                ]),
                (f"{spouse} — Elective Share Waiver (same appointment)", [
                    f"{spouse} signs the Waiver of Elective Share before the Notary",
                    f"{spouse} had opportunity to review with independent legal counsel",
                    "Notary notarizes signature on the Waiver as a separate document",
                    "Store Waiver separately from the Will (but keep them together)",
                ]),
            ]),
            ("PHASE 3 — AFTER SIGNING", [
                ("Document Storage", [
                    "Store original Will + Waiver in a fireproof home safe",
                    "DO NOT store the original Will in a safe deposit box (may require probate to access)",
                    f"Give {spouse} a copy; tell her where the original is stored",
                    "Give a copy to your estate planning attorney for their file",
                ]),
                ("Non-Probate Asset Updates  (complete within 30 days)", [
                    "Update 401k beneficiary designations per Fara'id percentages (see Document 4)",
                    "Update life insurance beneficiary designations",
                    "Update any POD/TOD bank or investment account designations",
                    "Re-title home deed from JTWROS to Tenants in Common if applicable",
                ]),
            ]),
            ("PHASE 4 — ONGOING MAINTENANCE", [
                ("Re-execute the Will after any of these events", [
                    "Death of a named heir (father, mother, or either daughter)",
                    f"Death or incapacity of {ep}, {ea}, {gp}, or {ga}",
                    "Birth of a new child",
                    "Marriage or divorce (Tennessee auto-revokes wills upon both)",
                    "Significant change in estate value (>25%)",
                    "Completion of the Mosul Waqf transfer — update or remove Article VI",
                ]),
                ("Routine", [
                    "Review the Will and all beneficiary designations every 3–5 years regardless of life events",
                ]),
            ]),
        ]

        for phase_title, sections in phases:
            article(doc, phase_title)
            for section_title, items in sections:
                body(doc, section_title, bold=True, space_before=8, space_after=2)
                for item in items:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.left_indent  = Inches(0.25)
                    p.paragraph_format.space_after  = Pt(2)
                    run = p.add_run(item)
                    _set_font(run)
            rule(doc)

        # Quick reference table
        article(doc, "TENNESSEE REQUIREMENTS — QUICK REFERENCE")
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _table_header(tbl, ["Requirement", "Detail", "Statute"], size=SMALL_PT)
        for row in [
            ("Writing", "Must be a written document", "T.C.A. § 32-1-104"),
            ("Testator Signature", "Must sign in presence of both witnesses", "T.C.A. § 32-1-104"),
            ("Two Witnesses", "Disinterested adults; sign in each other's presence", "T.C.A. § 32-1-104"),
            ("Self-Proving Affidavit", "Notarized at signing; eliminates witness testimony at probate", "T.C.A. § 32-1-104"),
            ("Spousal Elective Share", "Surviving spouse may claim 10–40% without waiver", "T.C.A. § 31-4-101"),
            ("Elective Share Waiver", "Spouse signs voluntarily with fair financial disclosure", "T.C.A. § 31-4-102"),
            ("Year's Support Waiver", "Included in Waiver document", "T.C.A. § 30-2-102"),
            ("Revocation — Marriage", "Prior will auto-revoked upon remarriage", "T.C.A. § 32-1-201"),
            ("Revocation — Divorce", "Ex-spouse provisions auto-revoked upon divorce", "T.C.A. § 32-1-202"),
            ("Guardianship", "Will nomination honored under 'best interest' standard", "T.C.A. § 36-6-106"),
            ("No-Contest Clause", "In terrorem clauses are enforceable", "T.C.A. § 32-3-119"),
        ]:
            _table_row(tbl, row, size=SMALL_PT)

        _disclaimer(doc)
        path = self.out / f"03_Execution_Checklist_{_safe(name)}.docx"
        doc.save(path)
        return path

    # ── DOCUMENT 4 — Beneficiary Designation Reminder ────────────────────────

    def gen_beneficiary(self) -> Path:
        doc    = _new_doc()
        name   = self._name()
        county = self._county()
        shares = self.dist.get("shares", [])
        net    = self.dist.get("net_mirath", 0)
        ret    = self.est.get("retirement_401k", 0)
        home_v = self.est.get("home", {}).get("value", 0)
        home_h = home_v / 2
        was_t  = self._wasiyyah_total()

        doc_title(doc, "BENEFICIARY DESIGNATION ACTION PLAN")
        sub_title(doc, f"{name}  ·  {county}")
        body(doc,
             "These assets pass OUTSIDE the Will via beneficiary designation. "
             "Update each account separately to enforce Fara'id distribution.",
             italic=True, space_after=12)
        rule(doc)

        # 401k
        if ret:
            article(doc, f"401(k) / RETIREMENT ACCOUNTS — Estimated Value: ${ret:,.0f}")
            body(doc,
                 "Action: Contact your plan administrator or HR department and submit "
                 "a new Beneficiary Designation form with the following allocations:")
            if shares:
                tbl = doc.add_table(rows=1, cols=4)
                tbl.style = "Table Grid"
                _table_header(tbl, ["Beneficiary", "Relationship", "Percentage", f"Est. Amount (${ret:,.0f})"])
                for s in shares:
                    _table_row(tbl, [
                        s.get("heir", ""),
                        s.get("relationship", ""),
                        f"{s.get('percentage',0):.2f}%",
                        f"${ret * s.get('percentage',0) / 100:,.0f}",
                    ])
                doc.add_paragraph()
            body(doc, "Plan Administrator: _______________________    Account #: _______________")
            rule(doc)

        # Home
        if home_v:
            article(doc, f"HOME — Estimated Value: ${home_v:,.0f} (Joint with Spouse)")
            body(doc,
                 "Pull your deed from the County Register of Deeds and confirm titling:")
            body(doc,
                 f"☐  Joint Tenants with Right of Survivorship (JTWROS): Full ${home_v:,.0f} "
                 "passes to spouse automatically — bypasses all other heirs. "
                 "Consult attorney about re-titling as Tenants in Common.",
                 indent=0.3)
            body(doc,
                 f"☐  Tenants in Common: Your 50% share (${home_h:,.0f}) passes through "
                 "this Will per Fara'id. No deed change needed.",
                 indent=0.3)
            rule(doc)

        # Full estate picture
        if shares:
            article(doc, "FULL ESTATE — COMPLETE ISLAMIC PICTURE")
            body(doc, "If all assets are updated to follow Fara'id:", italic=True)

            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Table Grid"
            _table_header(tbl,
                ["Heir", "Probate Will", "401k (update)", "Home 50% (re-title)", "TOTAL"],
                size=SMALL_PT)

            col_prob = col_ret = col_home = col_all = 0
            for s in shares:
                pct   = s.get("percentage", 0) / 100
                prob  = s.get("amount", 0)
                r_val = ret * pct
                h_val = home_h if s.get("relationship") == "Wife" else 0
                total = prob + r_val + h_val
                col_prob += prob; col_ret += r_val; col_home += h_val; col_all += total
                _table_row(tbl, [
                    s.get("heir", ""),
                    f"${prob:,.0f}",
                    f"${r_val:,.0f}",
                    f"${h_val:,.0f}" if h_val else "—",
                    f"${total:,.0f}",
                ], size=SMALL_PT)

            if was_t:
                _table_row(tbl, ["Wasiyyah (Waqf)", f"${was_t:,.0f}", "—", "—", f"${was_t:,.0f}"],
                           size=SMALL_PT)
                col_prob += was_t; col_all += was_t

            _table_row(tbl, [
                "TOTAL",
                f"${col_prob:,.0f}",
                f"${col_ret:,.0f}",
                f"${col_home:,.0f}",
                f"${col_all:,.0f}",
            ], bold=True, size=SMALL_PT)

            doc.add_paragraph()
            body(doc,
                 "* Spouse retains her own 50% of home as co-owner regardless of deed type.",
                 size=SMALL_PT, italic=True)
            rule(doc)

        _disclaimer(doc)
        path = self.out / f"04_Beneficiary_Designations_{_safe(name)}.docx"
        doc.save(path)
        return path


# ── CLI entry point ───────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate Islamic Will Word (.docx) documents from a JSON data file.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Example:\n"
            "  python script/generate_docs.py will_data.json\n"
            "  python script/generate_docs.py will_data.json --output-dir ./output/zaid\n\n"
            "The JSON file is produced by the /islamic-will skill (Stage 8).\n"
            "See script/will_data_template.json for the required schema."
        ),
    )
    parser.add_argument("data_file",     help="Path to will_data.json")
    parser.add_argument("--output-dir",  default="./output",
                        help="Directory to save .docx files (default: ./output)")
    args = parser.parse_args()

    data_path = Path(args.data_file)
    if not data_path.exists():
        print(f"Error: file not found — {data_path}", file=sys.stderr)
        sys.exit(1)

    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)

    out = Path(args.output_dir)
    print(f"\nIslamic Will Document Generator")
    print(f"{'─' * 40}")
    gen   = IslamicWillGenerator(data, out)
    paths = gen.generate_all()

    print(f"\nSaved to: {out.resolve()}")
    for p in paths:
        print(f"  ✓ {p.name}")
    print("\nNext: open each document, complete [BRACKETED FIELDS], then proceed to signing.")


if __name__ == "__main__":
    main()
